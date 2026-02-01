from flask import Flask, request, send_file
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Inches
import io


app = Flask(__name__)


# ================= PDF GENERATOR =================
def generate_pdf(data, image_file):

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)

    width, height = A4


    # ============ COLORED HEADER ============

    c.setFillColorRGB(0.1, 0.3, 0.6)
    c.rect(0, height - 140, width, 140, fill=1)

    header_y = height - 40


    # ---- Image (Circle) ----
    if image_file:

        img = Image.open(image_file).convert("RGBA")
        img = img.resize((100, 100))

        mask = Image.new("L", (100, 100), 0)
        draw = ImageDraw.Draw(mask)
        draw.ellipse((0, 0, 100, 100), fill=255)

        img.putalpha(mask)

        img_buffer = io.BytesIO()
        img.save(img_buffer, format="PNG")
        img_buffer.seek(0)

        c.drawImage(
            ImageReader(img_buffer),
            50,
            header_y - 100,
            width=100,
            height=100,
            mask="auto"
        )


    # White text
    c.setFillColorRGB(1, 1, 1)

    c.setFont("Helvetica-Bold", 22)
    c.drawString(170, header_y - 20, data["name"])

    c.setFont("Helvetica", 13)
    c.drawString(170, header_y - 45, data["job"])

    c.setFont("Helvetica", 11)
    c.drawString(
        170,
        header_y - 65,
        f'{data["email"]} | {data["phone"]}'
    )


    # Line
    c.setStrokeColorRGB(1, 0.5, 0)
    c.setLineWidth(2)
    c.line(50, height - 150, width - 50, height - 150)


    # Reset
    c.setFillColorRGB(0, 0, 0)
    c.setStrokeColorRGB(0, 0, 0)


    y = height - 190


    # ============ CONTENT ============

    sections = [
        ("About Me", data["about"]),
        ("Skills", data["skills"]),
        ("Experience", data["experience"]),
        ("Education", data["education"]),
    ]


    for title, content in sections:

        if y < 120:

            c.showPage()
            y = height - 60


        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, y, title)

        y -= 10
        c.line(50, y, width - 50, y)

        y -= 20
        c.setFont("Helvetica", 12)

        text = c.beginText(50, y)


        for line in content.split("\n"):

            if y < 100:

                c.drawText(text)
                c.showPage()

                text = c.beginText(50, height - 60)
                y = height - 60


            text.textLine(line)
            y -= 15


        c.drawText(text)

        y -= 40


    c.showPage()
    c.save()

    buffer.seek(0)

    return buffer



# ================= WORD GENERATOR =================
def generate_word(data, image_file):

    doc = Document()


    # Title
    doc.add_heading(data["name"], level=1)
    doc.add_paragraph(data["job"])
    doc.add_paragraph(f'{data["email"]} | {data["phone"]}')


    # Photo
    if image_file:

        image_file.seek(0)
        doc.add_picture(image_file, width=Inches(1.2))


    doc.add_paragraph("----------------------------")


    sections = [
        ("About Me", data["about"]),
        ("Skills", data["skills"]),
        ("Experience", data["experience"]),
        ("Education", data["education"]),
    ]


    for title, content in sections:

        doc.add_heading(title, level=2)
        doc.add_paragraph(content)


    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)

    return file_stream



# ================= MAIN =================
@app.route("/", methods=["GET", "POST"])
def home():

    if request.method == "POST":

        data = {
            "name": request.form.get("name"),
            "email": request.form.get("email"),
            "phone": request.form.get("phone"),
            "job": request.form.get("job"),
            "about": request.form.get("about"),
            "skills": request.form.get("skills"),
            "experience": request.form.get("experience"),
            "education": request.form.get("education"),
        }

        image = request.files.get("photo")

        action = request.form.get("action")


        if action == "pdf":

            file = generate_pdf(data, image)

            return send_file(
                file,
                as_attachment=True,
                download_name="cv_pro.pdf",
                mimetype="application/pdf"
            )


        if action == "word":

            file = generate_word(data, image)

            return send_file(
                file,
                as_attachment=True,
                download_name="cv_Rachid_pro.docx",
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )


    return """
    <html>
    <head>

        <title>CV Rachid Pro</title>

        <style>

            body {
                font-family: Arial;
                background: #1e1e1e;
                color: white;
            }

            form {
                width: 420px;
                margin: 40px auto;
                background: #2c2c2c;
                padding: 20px;
                border-radius: 10px;
            }

            input, textarea, button {
                width: 100%;
                margin: 6px 0;
                padding: 8px;
            }

            button {
                background: orange;
                border: none;
                font-weight: bold;
                cursor: pointer;
            }

            .pdf {
                background: #3498db;
            }

            .word {
                background: #2ecc71;
            }

            h2 {
                text-align: center;
            }

        </style>

    </head>

    <body>

        <form method="POST" enctype="multipart/form-data">

            <h2>CV Rachid Pro 📄📝</h2>

            <input name="name" placeholder="Full Name" required>
            <input name="email" placeholder="Email" required>
            <input name="phone" placeholder="Phone" required>
            <input name="job" placeholder="Job" required>

            <label>Photo</label>
            <input type="file" name="photo" accept="image/*" required>

            <textarea name="about" placeholder="About you"></textarea>
            <textarea name="skills" placeholder="Skills"></textarea>
            <textarea name="experience" placeholder="Experience"></textarea>
            <textarea name="education" placeholder="Education"></textarea>


            <button class="pdf" name="action" value="pdf">Download PDF</button>

            <button class="word" name="action" value="word">Download Word</button>

        </form>

    </body>
    </html>
    """



# ================= RUN =================
if __name__ == "__main__":
    app.run(debug=True, port=5004)
